from docx import Document as DocxDocument
from pathlib import Path
from config import Settings

def save_docx(md_text: str, settings: Settings) -> str:
    doc: DocxDocument = DocxDocument()
    for line in md_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        else:
            doc.add_paragraph(line)

    path = Path(settings.OUTPUT_DIR) / "review.docx"
    doc.save(path)
    return str(path)
