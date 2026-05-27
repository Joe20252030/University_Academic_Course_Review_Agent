from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt

from uacragent.export._utils import safe_timestamp
from uacragent.infra.workspace import WorkspacePaths

# Regex that splits a text span into inline formatting groups.
# Groups (matched in priority order — longest/most-specific first):
#   1,2  — ***bold-italic***  (triple asterisks)
#   3,4  — **bold** or __bold__
#   5,6  — *italic* or _italic_
#   7    — `code` (backtick code span)
_INLINE_RE = re.compile(
    r"(\*\*\*)(.+?)\1"    # bold+italic: ***text***
    r"|"
    r"(\*\*|__)(.+?)\3"   # bold:        **text** or __text__
    r"|"
    r"(\*|_)(.+?)\5"      # italic:      *text*  or _text_
    r"|"
    r"`(.+?)`",            # code span:   `text`
    re.DOTALL,
)


def _add_inline_runs(paragraph, text: str, base_size_pt: int = 11) -> None:
    """Add runs to *paragraph* with inline formatting parsed from *text*.

    Handles:
    - ``***bold-italic***`` — bold + italic (matched before **bold** and *italic*)
    - ``**bold**`` / ``__bold__`` — bold
    - ``*italic*`` / ``_italic_`` — italic
    - `` `code` `` — backtick code span (monospace, no other styling)

    Plain text segments between markers are added as unstyled runs.
    All runs are set to *base_size_pt* so body paragraph sizing is preserved.
    """
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            run.font.size = Pt(base_size_pt)

        if m.group(1):      # bold+italic marker (***)
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):    # bold marker (** or __)
            run = paragraph.add_run(m.group(4))
            run.bold = True
        elif m.group(5):    # italic marker (* or _)
            run = paragraph.add_run(m.group(6))
            run.italic = True
        else:               # code span (`code`)
            run = paragraph.add_run(m.group(7))
            # Apply monospace font for code spans; fall back to plain text if
            # Courier New is not available (docx will use the default mono font).
            try:
                run.font.name = "Courier New"
            except Exception:  # noqa: BLE001
                pass
        run.font.size = Pt(base_size_pt)
        pos = m.end()

    # Trailing plain text
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(base_size_pt)


def _add_markdown_to_docx(doc: DocxDocument, md_text: str) -> None:
    """Convert simplified Markdown into python-docx paragraphs.

    Handles:
    - Headings (# / ## / ###)
    - Bullet lists  (- / *)
    - Numbered lists (1. / 2. …)
    - Blank lines preserved as empty paragraphs (spacing)
    - Inline **bold** / *italic* / __bold__ / _italic_ within body text

    This is intentionally simple; a full Markdown parser is overkill for
    review documents, but inline bold/italic is critical for academic content.
    """
    for line in md_text.splitlines():
        stripped = line.strip()

        # Blank line → empty paragraph preserves paragraph spacing
        if not stripped:
            doc.add_paragraph("")
            continue

        is_heading = False

        # Headings — let the built-in heading style control the font size
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            is_heading = True
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            is_heading = True
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            is_heading = True
        # Bullet lists
        elif re.match(r"^[-*]\s", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, stripped[2:])
        # Numbered lists
        elif re.match(r"^\d+\.\s", stripped):
            body = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, body)
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, stripped)

        # Headings inherit size from their style; no per-run override needed.
        _ = is_heading  # referenced to suppress linter "assigned but unused"


def save_docx(md_text: str, workspace_paths: WorkspacePaths) -> str:
    Path(workspace_paths.outputs).mkdir(parents=True, exist_ok=True)

    doc = DocxDocument()
    _add_markdown_to_docx(doc, md_text)

    path = Path(workspace_paths.outputs) / f"review_{safe_timestamp()}.docx"
    doc.save(str(path))
    return str(path)
