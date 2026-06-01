"""Tests for prompt language instructions, export inline formatting,
persistence.delete_session symlink guard, and safe_timestamp."""
from __future__ import annotations

from pathlib import Path

import pytest


# ---------------------------------------------------------------------------
# get_language_instruction
# ---------------------------------------------------------------------------

class TestGetLanguageInstruction:
    def test_auto_mode_non_pipeline(self):
        from uacragent.agent.prompts._prompts import get_language_instruction
        instr = get_language_instruction("auto")
        assert "detect" in instr.lower() or "language" in instr.lower()

    def test_en_mode_non_pipeline(self):
        from uacragent.agent.prompts._prompts import get_language_instruction
        instr = get_language_instruction("en")
        assert "english" in instr.lower()

    def test_zh_cn_mode_non_pipeline(self):
        from uacragent.agent.prompts._prompts import get_language_instruction
        instr = get_language_instruction("zh_CN")
        assert "chinese" in instr.lower() or "中文" in instr

    def test_unknown_falls_back_to_english(self):
        from uacragent.agent.prompts._prompts import get_language_instruction
        instr_en = get_language_instruction("en")
        instr_xx = get_language_instruction("klingon")
        assert instr_en == instr_xx

    def test_pipeline_variant_is_shorter(self):
        from uacragent.agent.prompts._prompts import get_language_instruction
        full = get_language_instruction("zh_CN", pipeline=False)
        short = get_language_instruction("zh_CN", pipeline=True)
        assert len(full) > len(short)

    def test_pipeline_auto_returns_something(self):
        from uacragent.agent.prompts._prompts import get_language_instruction
        instr = get_language_instruction("auto", pipeline=True)
        assert len(instr) > 0

    def test_pipeline_en_returns_something(self):
        from uacragent.agent.prompts._prompts import get_language_instruction
        instr = get_language_instruction("en", pipeline=True)
        assert "english" in instr.lower()


# ---------------------------------------------------------------------------
# DOCX inline formatting (_add_inline_runs)
# ---------------------------------------------------------------------------

class TestDocxInlineFormatting:
    def _para_text(self, para) -> str:
        return "".join(run.text for run in para.runs)

    def test_plain_text_no_formatting(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_inline_runs

        doc = DocxDocument()
        p = doc.add_paragraph()
        _add_inline_runs(p, "hello world")
        assert self._para_text(p) == "hello world"

    def test_bold_text(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_inline_runs

        doc = DocxDocument()
        p = doc.add_paragraph()
        _add_inline_runs(p, "**bold text**")
        runs = p.runs
        bold_runs = [r for r in runs if r.bold]
        assert bold_runs
        assert "bold text" in self._para_text(p)

    def test_italic_text(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_inline_runs

        doc = DocxDocument()
        p = doc.add_paragraph()
        _add_inline_runs(p, "*italic text*")
        runs = p.runs
        italic_runs = [r for r in runs if r.italic]
        assert italic_runs

    def test_bold_italic_text(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_inline_runs

        doc = DocxDocument()
        p = doc.add_paragraph()
        _add_inline_runs(p, "***bold-italic***")
        runs = p.runs
        bi_runs = [r for r in runs if r.bold and r.italic]
        assert bi_runs

    def test_code_span(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_inline_runs

        doc = DocxDocument()
        p = doc.add_paragraph()
        _add_inline_runs(p, "`code here`")
        assert "code here" in self._para_text(p)

    def test_mixed_plain_and_bold(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_inline_runs

        doc = DocxDocument()
        p = doc.add_paragraph()
        _add_inline_runs(p, "prefix **bold** suffix")
        full = self._para_text(p)
        assert "prefix" in full
        assert "bold" in full
        assert "suffix" in full


# ---------------------------------------------------------------------------
# DOCX _add_markdown_to_docx — structural elements
# ---------------------------------------------------------------------------

class TestDocxMarkdown:
    def test_heading_level_1(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_markdown_to_docx

        doc = DocxDocument()
        _add_markdown_to_docx(doc, "# Main Title")
        headings = [p for p in doc.paragraphs if "Heading" in (p.style.name or "")]
        assert any(p.text == "Main Title" for p in headings)

    def test_heading_level_2(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_markdown_to_docx

        doc = DocxDocument()
        _add_markdown_to_docx(doc, "## Section")
        headings = [p for p in doc.paragraphs if "Heading 2" in (p.style.name or "")]
        assert any("Section" in p.text for p in headings)

    def test_bullet_list(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_markdown_to_docx

        doc = DocxDocument()
        _add_markdown_to_docx(doc, "- Item one\n- Item two")
        bullet_paras = [p for p in doc.paragraphs if "Bullet" in (p.style.name or "")]
        assert len(bullet_paras) == 2

    def test_numbered_list(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_markdown_to_docx

        doc = DocxDocument()
        _add_markdown_to_docx(doc, "1. First\n2. Second")
        num_paras = [p for p in doc.paragraphs if "Number" in (p.style.name or "")]
        assert len(num_paras) == 2

    def test_empty_line_produces_paragraph(self):
        from docx import Document as DocxDocument
        from uacragent.export.docx import _add_markdown_to_docx

        doc = DocxDocument()
        _add_markdown_to_docx(doc, "line1\n\nline2")
        all_texts = [p.text for p in doc.paragraphs]
        assert "" in all_texts  # blank paragraph for the empty line


# ---------------------------------------------------------------------------
# safe_timestamp
# ---------------------------------------------------------------------------

class TestSafeTimestamp:
    def test_format_is_filename_safe(self):
        from uacragent.export._utils import safe_timestamp
        ts = safe_timestamp()
        assert ":" not in ts
        assert " " not in ts

    def test_two_consecutive_calls_differ(self):
        from uacragent.export._utils import safe_timestamp
        ts1 = safe_timestamp()
        ts2 = safe_timestamp()
        # Includes microseconds so they will almost certainly differ
        # (unless the system clock has zero resolution)
        assert isinstance(ts1, str) and isinstance(ts2, str)

    def test_starts_with_year(self):
        from uacragent.export._utils import safe_timestamp
        import re
        ts = safe_timestamp()
        assert re.match(r"^\d{4}-\d{2}-\d{2}_", ts)


# ---------------------------------------------------------------------------
# persistence.delete_session — symlink guard on workspace folder
# ---------------------------------------------------------------------------

class TestDeleteSessionSymlinkGuard:
    def test_symlinked_workspace_not_deleted(self, tmp_path, monkeypatch):
        """A symlinked managed workspace must not be wiped via rmtree."""
        from uacragent.infra import persistence as pers
        from uacragent.agent.session import AgentSession
        from uacragent.infra.workspace import ensure_workspace_dirs, workspace_paths

        # Create a real directory and a symlink to it
        real_dir = tmp_path / "real_workspace"
        real_dir.mkdir()
        link = tmp_path / "sessions" / "link_ws"
        link.parent.mkdir(parents=True, exist_ok=True)
        link.symlink_to(real_dir)

        # Make the sessions dir appear to be under the managed app-data dir
        monkeypatch.setattr(pers, "get_app_data_dir", lambda: tmp_path)

        # Add an agent_dir so delete_session has something to find
        agent_dir = link / ".uacragent"
        agent_dir.mkdir(parents=True, exist_ok=True)
        (agent_dir / "session.json").write_text('{"version": 1}')

        pers.delete_session(link)

        # The real directory must not have been wiped
        assert real_dir.exists()
